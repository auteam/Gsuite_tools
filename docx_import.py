import docx
from docx.api import Document
import re
from transliterate import translit


class ImportDoc:
    filename = 'source/КСи Д/Д-115.docx'
    groups = ['bi-15']
    grouplist = []
    text = ''
    pattern = 'n1'

    def __init__(self, filename, pattern):
        self.filename = filename
        self.pattern = pattern

        if pattern == 'n1':
            self.grouplist = self.get_grouplist(filename, pattern)
            self.groups = self.normalise_group(self.get_text(filename), pattern)
        elif pattern == 'n2':
            self.grouplist = self.get_grouplist(filename, pattern)
            self.groups = self.normalise_group(self.get_text(filename), pattern)
        # elif pattern == 'n3':
        #     self.grouplist = self.get_grouplist(filename, pattern)
        #     self.groups = self.normalise_group(self.get_text(filename), pattern)

    def get_text(self, filename):
        doc = docx.Document(filename)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text.split('\n'))
        return full_text

    def normalise_group(self, p, doc_pattern):
        def normalise(group):
            if len(group) == 5:  # для групп с 1 буквой. Л Д Э
                group = group[0:2] + group[1 - 3:]
                return group
            elif len(group) == 6:  # для групп с 2мя буквами КС Би
                group = group[0:3] + group[1 - 3:]
                return group
            elif len(group) == 7:  # заочники с 3мя буквами НЕ ПРОВЕРЕНО
                group = group[0:4] + group[1 - 4:]
                return group
            else:
                print('INCORRECT GROUP NAME')

        if doc_pattern == 'n1':
            groups = []
            group_pattern = r'\w+\-([0-9])([0-9])([0-9])'
            for row in p:
                if row[0] is not None:
                    match = re.search(group_pattern, row[0])
                    if match:
                        group_ru = match.group()
                        group = (translit(group_ru, 'ru', reversed=True)).lower()
                        groups.append(normalise(group))
            return groups

        elif doc_pattern == 'n2':
            group_pattern = r'\w+\-([0-9])([0-9])([0-9])*'
            groups = []
            for doc in p:
                for string in doc:
                    if re.match(group_pattern, string):
                        group_ru = re.match(group_pattern, string).group(0)
                        group = (translit(group_ru, 'ru', reversed=True)).lower()
                        groups.append(normalise(group))
            return groups

        # elif doc_pattern == 'n3':
        #     group_pattern = r"*-*"
        #     groups = []
        #
        #     for i in range(len(p)):
        #         if p[i][0] is not None:
        #             pass
        #             # print(p[i][0])

    def get_grouplist(self, filename, pattern):
        def none_test(names):
            result = 0
            for name in names:
                if not bool(name):
                    result += 1
            if result == 0:
                return True
            else:
                return False

        if pattern == 'n1':
            tables = Document(filename).tables
            groups_list = []  # list of group lists

            for i in range(len(tables)):
                group_list = []
                for j in range(len(tables[i].columns)):  # Получаем cells из column ФИО
                    for cell in tables[i].columns[j].cells:
                        for paragraph in cell.paragraphs:
                            group_list.append(paragraph.text.split(' '))   # Разделение ФИО

                for k in range(len(group_list)):
                    if len(group_list[k]) == 2 or len(group_list[k]) == 3:
                        if none_test(group_list[k]):
                            groups_list.append(group_list[k])
                return groups_list

        elif pattern == 'n2':
            tables = Document(filename).tables
            groups_list = []  # list of group lists

            for i in range(len(tables)):
                group_list = []
                for j in range(1, len(tables[i].columns), 2):  # Получаем cells из column ФИО
                    for cell in tables[i].columns[j].cells:
                        for paragraph in cell.paragraphs:
                            group_list.append(paragraph.text.split(','))

                groups_list.append(group_list)

            return groups_list

        # elif pattern == 'n3':
        #     text = self.get_text(filename)
        #     group_pattern = r''
        #
        #     text_lists = []
        #     for row in text:
        #         text_lists.append(row[0].split(' '))
        #     text_lists = [x for x in text_lists if x[0] is not None]
        #     groups_list = [user for user in text_lists if user]  # list of group lists
        else:
            print("WRONG DOC PATTERN")
