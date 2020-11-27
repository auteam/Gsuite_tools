import docx
import csv
from docx.api import Document
import re
from transliterate import translit
from config import config
from src.student import *
from src.docx_import import *


class ImportDoc:
    filename = 'source/КСи Д/Д-115.docx'
    group = 'bi-15'
    grouplist = []
    text = ''
    pattern = 'n1'

    def __init__(self, filename, pattern, split_fio):
        self.filename = filename
        self.pattern = pattern
        # TODO отдельные классы для импорта docx и csv
        if pattern == 'n1':
            self.grouplist = self.get_grouplist(filename)
            self.group = self.normalise_group(self.get_text(filename))
        if pattern == 'csv':
            self.grouplist = self.get_fio_csv(filename, bool(split_fio))
            self.group = ''

    def get_fio_csv(self, filename, split_fio):
        users = []
        with open(filename) as csvfile:
            spamreader = csv.reader(csvfile, delimiter=',')
            for row in spamreader:
                if bool(row):                  # ignore empty lines
                    if config.import_conf.split_FIO:
                        users.append(row[0])
                    else:
                        users.append(row[0:3])
        print(users)
        return users

    def get_text(self, filename):
        doc = docx.Document(filename)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        return ''.join(full_text)

    def normalise_group(self, p):
        group_pattern = r'\w+\-([0-9])([0-9])([0-9])$'
        match = re.search(group_pattern, p)

        group_ru = match.group()

        group = (translit(group_ru, 'ru', reversed=True)).lower()
        group = group[0:2] + group[1 - 3:]
        return group

    def get_grouplist(self, filename):
        doc = Document(filename)
        table = doc.tables[0]

        # Data will be a list of rows represented as dictionaries
        # containing each row's data.
        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)

        grouplist = []

        for dict_ in data:
            fio = dict_.get('Фамилия, имя, отчество ')
            grouplist.append(fio)

        return grouplist
