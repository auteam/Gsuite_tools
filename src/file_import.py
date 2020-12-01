from config import config

import csv
import sys

import docx
from docx.api import Document
import re
from transliterate import translit


# TODO CheckFile class, to check incoming files
# TODO Group handling
class ImportFile:
    """
    filename = 'source/КСи Д/Д-115.docx'
    grouplist = []
    """

    def __init__(self, filename):
        self.filename = filename
        pattern = config.import_conf.pattern

        if pattern == 'n1':
            docx_import = self.ImportDocx(self.filename)
            self.grouplist = docx_import.get_grouplist(filename)
        if pattern == 'csv':
            csv_import = self.ImportCSV(self.filename)
            self.grouplist = csv_import.grouplist

    class ImportCSV:

        def __init__(self, filename):
            self.grouplist = self.get_fio_csv(filename)

        def get_fio_csv(self, filename):
            users = []
            try:
                with open(filename, encoding=config.import_conf.encoding) as csvfile:
                    spamreader = csv.reader(csvfile, delimiter=',')
                    for row in spamreader:
                        if bool(row):  # ignore empty lines
                            if config.import_conf.split_FIO:
                                users.append(row[0])
                            else:
                                users.append(row[0:3])
            except:
                print("ERROR:", sys.exc_info()[0])
            print(users)
            return users

    class ImportDocx:

        def __init__(self, filename):
            self.grouplist = []

        def get_text(self, filename):
            doc = docx.Document(filename)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            return ''.join(full_text)

        def normalise_group(self, p):
            group_pattern = r'\w+\-([0-9])([0-9])([0-9])$'
            match = re.search(group_pattern, p)

            group_ru = match.group()

            group = (translit(group_ru, 'ru', reversed=True)).lower()
            group = group[0:2] + group[1 - 3:]
            return group

        def get_grouplist(self, filename):
            doc = Document(filename)
            table = doc.tables[0]

            # Data will be a list of rows represented as dictionaries
            # containing each row's data.
            data = []

            keys = None
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)

                # Establish the mapping based on the first row
                # headers; these will become the keys of our dictionary
                if i == 0:
                    keys = tuple(text)
                    continue

                # Construct a dictionary for this row, mapping
                # keys to values for this row
                row_data = dict(zip(keys, text))
                data.append(row_data)

            grouplist = []

            for dict_ in data:
                fio = dict_.get('Фамилия, имя, отчество ')
                grouplist.append(fio)

            return grouplist
